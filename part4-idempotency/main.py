import os
import re
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
import sys
import stat
import hashlib

# Load environment variables

load_dotenv()

SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
GOOGLE_SHEETS_KEY = os.getenv("GOOGLE_SHEETS_KEY")

# Fail-fast environment validation

if not all([SPREADSHEET_ID, GOOGLE_SHEETS_KEY]):
    print("ERROR: Missing required environment variables.", file=sys.stderr)
    print(
        "  Set SPREADSHEET_ID and GOOGLE_SHEETS_KEY"
        " in your .env file.", file=sys.stderr)
    sys.exit(1)

# Output directory provisioning

BASE_OUTPUT_DIR = os.path.join(os.getcwd(), "submissions", "customers")
os.makedirs(BASE_OUTPUT_DIR, exist_ok=True)

if not os.access(BASE_OUTPUT_DIR, os.W_OK):
    print(
        f"ERROR: Output directory is not writable:"
        f" {BASE_OUTPUT_DIR}", file=sys.stderr)
    sys.exit(1)

# Markers directory provisioning

MARKERS_DIR = os.path.join(os.getcwd(), "markers")
os.makedirs(MARKERS_DIR, mode=0o700, exist_ok=True)

if not os.access(MARKERS_DIR, os.W_OK):
    print(
        f"ERROR: Markers directory is not writable:"
        f" {MARKERS_DIR}", file=sys.stderr)
    sys.exit(1)

# Enforce key permissions

def enforce_key_permission(path: str) -> None:
    mode = stat.S_IMODE(os.stat(path).st_mode)
    if mode != 0o600:
        os.chmod(path, 0o600)
        mode = stat.S_IMODE(os.stat(path).st_mode)
        if mode != 0o600:
            print(
                f"ERROR: Could not enforce permissions on {path}",
                file=sys.stderr)
            sys.exit(1)

# Service account key validation

if not os.path.isfile(GOOGLE_SHEETS_KEY):
    print(f"ERROR: Key file not found: {GOOGLE_SHEETS_KEY}", file=sys.stderr)
    sys.exit(1)

enforce_key_permission(GOOGLE_SHEETS_KEY)

# Helper function to convert a customer name into a safe folder name

def normalise_name(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^\w\s-]", "", name)
    name = re.sub(r"\s+", "_", name)
    return name

# Path confinement

def _safe_join(base_dir: str, untrusted_segment: str) -> str:
    resolved_base = os.path.realpath(base_dir)
    candidate = os.path.realpath(os.path.join(base_dir, untrusted_segment))
    if (not candidate.startswith(resolved_base + os.sep)
            and candidate != resolved_base):
        print(
            f"ERROR: Path traversal detected:"
            f" '{untrusted_segment}' escapes sandbox",
            file=sys.stderr)
        sys.exit(1)
    return candidate

# Content-derived identifier for deduplication

def fingerprint(row: list) -> str:
    joined = "\x00".join(cell.strip() for cell in row)
    return hashlib.sha256(joined.encode()).hexdigest()

# Atomically create a marker file; succeeds only once

def claim_marker(marker_file: str) -> bool:
    try:
        fd = os.open(marker_file, os.O_CREAT | os.O_EXCL, 0o600)
        os.close(fd)
        return True
    except FileExistsError:
        return False

# Authenticate with Google Sheets

credentials = service_account.Credentials.from_service_account_file(
    GOOGLE_SHEETS_KEY,
    scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"],
)

sheets_service = build("sheets", "v4", credentials=credentials)

# Read spreadsheet data

sheet_range = "Form Responses 1!A:Z"

response = sheets_service.spreadsheets().values().get(
    spreadsheetId=SPREADSHEET_ID,
    range=sheet_range,
).execute()

values = response.get("values", [])

headers = values[0]
rows = values[1:]

name_index = headers.index("Name")

# Process each form submission

processed = 0
skipped = 0

for row_number, row in enumerate(rows, start=1):
    customer_name = row[name_index].strip()

    if not customer_name:
        skipped += 1
        continue

    folder_name = normalise_name(customer_name)
    customer_folder = _safe_join(BASE_OUTPUT_DIR, folder_name)

    # Idempotency gate

    row_id = fingerprint(row)
    marker_file = os.path.join(MARKERS_DIR, f"{folder_name}__{row_id}.done")

    if not claim_marker(marker_file):
        print(f"Skipping row {row_number} (already processed): {customer_name}")
        skipped += 1
        continue

    os.makedirs(customer_folder, exist_ok=True)

    document = Document()
    document.add_heading(f"Customer Submission - {customer_name}", level=1)

    for header, cell in zip(headers, row):
        document.add_paragraph(f"{header}: {cell}")

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S_%f")
    filename = f"{timestamp}_submission.docx"
    file_path = os.path.join(customer_folder, filename)

    document.save(file_path)

    processed += 1
    print(f"Saved submission {row_number}: {file_path}")

print(f"Completed: {processed} processed, {skipped} skipped.")
print(f"To reprocess a row, delete its .done file in {MARKERS_DIR}")