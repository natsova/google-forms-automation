import os
import re
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
import sys
import stat
import hashlib
import glob
import logging

# Load environment variables

load_dotenv()

SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
GOOGLE_SHEETS_KEY = os.getenv("GOOGLE_SHEETS_KEY")

# Logging setup

os.makedirs("logs", mode=0o700, exist_ok=True)
LOG_FILE = os.path.join("logs", "forms_automation.log")
_log_fd = os.open(LOG_FILE, os.O_WRONLY | os.O_CREAT | os.O_APPEND, 0o600)
os.close(_log_fd)

logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(message)s",
    datefmt="%H:%M:%S",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(LOG_FILE, mode='a')
    ]
)
log = logging.getLogger("forms-automation")

# Fail-fast environment validation

if not all([SPREADSHEET_ID, GOOGLE_SHEETS_KEY]):
    log.error("Missing required environment variables.")
    log.error("  Set SPREADSHEET_ID and GOOGLE_SHEETS_KEY in your .env file.")
    sys.exit(1)

# Enforce key permissions

def enforce_key_permission(path: str) -> None:
    mode = stat.S_IMODE(os.stat(path).st_mode)
    if mode != 0o600:
        os.chmod(path, 0o600)
        mode = stat.S_IMODE(os.stat(path).st_mode)
        if mode != 0o600:
            log.error("Could not enforce permissions on %s", path)
            sys.exit(1)

# Service account key validation

if not os.path.isfile(GOOGLE_SHEETS_KEY):
    log.error("Key file not found: %s", GOOGLE_SHEETS_KEY)
    sys.exit(1)

enforce_key_permission(GOOGLE_SHEETS_KEY)

# Output directory provisioning

BASE_OUTPUT_DIR = os.path.join(os.getcwd(), "submissions", "customers")
os.makedirs(BASE_OUTPUT_DIR, exist_ok=True)

if not os.access(BASE_OUTPUT_DIR, os.W_OK):
    log.error("Output directory is not writable: %s", BASE_OUTPUT_DIR)
    sys.exit(1)

# Markers directory provisioning

MARKERS_DIR = os.path.join(os.getcwd(), "markers")
os.makedirs(MARKERS_DIR, mode=0o700, exist_ok=True)

if not os.access(MARKERS_DIR, os.W_OK):
    log.error("Markers directory is not writable: %s", MARKERS_DIR)
    sys.exit(1)

# Orphan cleanup

for orphan in glob.glob(
        os.path.join(BASE_OUTPUT_DIR, "**", "*.tmp"), recursive=True):
    os.remove(orphan)
    log.info("Removed orphaned temp file: %s", orphan)

# Helper function to convert a customer name into a safe folder name

def normalise_name(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^\w\s-]", "", name)
    name = re.sub(r"\s+", "_", name)
    return name

# Path confinement

def _safe_join(base_dir: str, untrusted_segment: str) -> str:
    resolved_base = os.path.realpath(base_dir)
    candidate = os.path.realpath(os.path.join(base_dir, untrusted_segment))
    if (not candidate.startswith(resolved_base + os.sep)
            and candidate != resolved_base):
        log.error(
            "Path traversal detected: '%s' escapes sandbox",
            untrusted_segment)
        sys.exit(1)
    return candidate

# Content-derived identifier for deduplication

def fingerprint(row: list) -> str:
    joined = "\x00".join(cell.strip() for cell in row)
    return hashlib.sha256(joined.encode()).hexdigest()

# Atomically create a marker file; succeeds only once

def claim_marker(marker_file: str) -> bool:
    try:
        fd = os.open(marker_file, os.O_CREAT | os.O_EXCL, 0o600)
        os.close(fd)
        return True
    except FileExistsError:
        return False

# Authenticate with Google Sheets

credentials = service_account.Credentials.from_service_account_file(
    GOOGLE_SHEETS_KEY,
    scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"],
)

sheets_service = build("sheets", "v4", credentials=credentials)

log.info("Authenticated with Google Sheets")

# Read spreadsheet data

sheet_range = "Form Responses 1!A:Z"

response = sheets_service.spreadsheets().values().get(
    spreadsheetId=SPREADSHEET_ID,
    range=sheet_range,
).execute()

values = response.get("values", [])

headers = values[0]
rows = values[1:]

name_index = headers.index("Name")

log.info("Fetched %d row(s) from spreadsheet", len(rows))

# Process each form submission

processed = 0
skipped = 0

for row_number, row in enumerate(rows, start=1):
    customer_name = row[name_index].strip()

    if not customer_name:
        skipped += 1
        continue

    folder_name = normalise_name(customer_name)
    customer_folder = _safe_join(BASE_OUTPUT_DIR, folder_name)

    # Idempotency gate

    row_id = fingerprint(row)
    marker_file = os.path.join(MARKERS_DIR, f"{folder_name}__{row_id}.done")

    if not claim_marker(marker_file):
        log.info("Row %d already processed, skipping: %s", row_number, customer_name)
        skipped += 1
        continue

    os.makedirs(customer_folder, exist_ok=True)

    document = Document()
    document.add_heading(f"Customer Submission - {customer_name}", level=1)

    for header, cell in zip(headers, row):
        document.add_paragraph(f"{header}: {cell}")

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S_%f")
    filename = f"{timestamp}_submission.docx"
    file_path = os.path.join(customer_folder, filename)

    # Atomic promotion with rollback

    tmp_file = file_path + ".tmp"
    try:
        document.save(tmp_file)
        os.replace(tmp_file, file_path)
    except Exception:
        for path in (marker_file, tmp_file):
            if os.path.exists(path):
                os.unlink(path)
        log.exception("Row %d failed, rolled back: %s", row_number, customer_name)
        raise

    processed += 1
    log.info("Row %d -> %s", row_number, file_path)

log.info("Completed: %d processed, %d skipped", processed, skipped)

if processed == 0 and skipped > 0:
    log.info(
        "All rows were already processed. To force a reprocess, delete the "
        "corresponding .done file(s) in %s", MARKERS_DIR)
else:
    log.info("To reprocess a row, delete its .done file in %s", MARKERS_DIR)
