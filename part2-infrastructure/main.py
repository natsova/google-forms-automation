import os
import re
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
import sys

# Load environment variables

load_dotenv()

SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
GOOGLE_SHEETS_KEY = os.getenv("GOOGLE_SHEETS_KEY")
BASE_OUTPUT_DIR = os.path.join(os.getcwd(), "submissions", "customers")

# Fail-fast environment validation
if not all([SPREADSHEET_ID, GOOGLE_SHEETS_KEY]):
    print("ERROR: Missing required environment variables.", file=sys.stderr)
    print(
        "  Set SPREADSHEET_ID and GOOGLE_SHEETS_KEY"
        " in your .env file.", file=sys.stderr)
    sys.exit(1)

# Service account key validation
if not os.path.isfile(GOOGLE_SHEETS_KEY):
    print(f"ERROR: Key file not found: {GOOGLE_SHEETS_KEY}", file=sys.stderr)
    sys.exit(1)

# Output directory provisioning
BASE_OUTPUT_DIR = os.path.abspath(BASE_OUTPUT_DIR)
os.makedirs(BASE_OUTPUT_DIR, exist_ok=True)

if not os.access(BASE_OUTPUT_DIR, os.W_OK):
    print(
        f"ERROR: Output directory is not writable:"
        f" {BASE_OUTPUT_DIR}", file=sys.stderr)
    sys.exit(1)

# Helper function to convert a customer name into a safe folder name

def normalise_name(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^\w\s-]", "", name)
    name = re.sub(r"\s+", "_", name)
    return name

# Authenticate with Google Sheets

credentials = service_account.Credentials.from_service_account_file(
    GOOGLE_SHEETS_KEY,
    scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"],
)

sheets_service = build("sheets", "v4", credentials=credentials)

# Read spreadsheet data

sheet_range = "Form Responses 1!A:Z"

response = sheets_service.spreadsheets().values().get(
    spreadsheetId=SPREADSHEET_ID,
    range=sheet_range,
).execute()

values = response.get("values", [])

headers = values[0]
rows = values[1:]

name_index = headers.index("Name")

# Process each form submission

for row_number, row in enumerate(rows, start=1):
    customer_name = row[name_index].strip()

    if not customer_name:
        continue

    folder_name = normalise_name(customer_name)
    customer_folder = os.path.join(BASE_OUTPUT_DIR, folder_name)
    os.makedirs(customer_folder, exist_ok=True)

    document = Document()
    document.add_heading(f"Customer Submission - {customer_name}", level=1)

    for header, cell in zip(headers, row):
        document.add_paragraph(f"{header}: {cell}")

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S_%f")
    filename = f"{timestamp}_submission.docx"
    file_path = os.path.join(customer_folder, filename)

    document.save(file_path)

    print(f"Saved submission {row_number}: {file_path}")

print("All submissions processed.")

